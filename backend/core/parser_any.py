# core/parser_any.py
from __future__ import annotations
import io
from typing import Optional

import filetype
import chardet

# PDF
import fitz
import pdfplumber

# DOCX
from docx import Document

# HTML
from bs4 import BeautifulSoup

from .layout_pdf import extract_pdf_text_layout
from .text_cleaner import clean_text, drop_repeated_header_footer


def _pdf_text_pymupdf(pdf_bytes: bytes) -> str:
    text = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            text.append(page.get_text("text"))
    return "\n".join(text)


def _pdf_text_pdfplumber(pdf_bytes: bytes) -> str:
    text = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            try:
                text.append(page.extract_text() or "")
            except Exception:
                pass
    return "\n".join(text)


def extract_pdf_text_any(pdf_bytes: bytes) -> str:
    """
    Order of attempts (no OCR):
      1) Layout-aware via PyMuPDF blocks (handles sidebars/2 columns)
      2) Native PyMuPDF text
      3) pdfplumber text
    """
    # 1) layout-aware
    try:
        t = extract_pdf_text_layout(pdf_bytes)
        if t.strip():
            return t
    except Exception:
        pass

    # 2) native
    try:
        t = _pdf_text_pymupdf(pdf_bytes)
        if t.strip():
            return t
    except Exception:
        pass

    # 3) fallback
    try:
        t = _pdf_text_pdfplumber(pdf_bytes)
        if t.strip():
            return t
    except Exception:
        pass

    return ""


def extract_docx_text(docx_bytes: bytes) -> str:
    bio = io.BytesIO(docx_bytes)
    doc = Document(bio)
    chunks = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            chunks.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def extract_html_text(html_bytes: bytes) -> str:
    enc = chardet.detect(html_bytes).get("encoding") or "utf-8"
    html = html_bytes.decode(enc, errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()
    return soup.get_text(separator="\n")


def extract_textlike_text(raw_bytes: bytes) -> str:
    enc = chardet.detect(raw_bytes).get("encoding") or "utf-8"
    return raw_bytes.decode(enc, errors="ignore")


def _guess_kind(file_bytes: bytes, filename: Optional[str], content_type: Optional[str]) -> str:
    kind = filetype.guess(file_bytes)
    if kind:
        if kind.mime == "application/pdf":
            return "pdf"
        if kind.mime in ("text/html",):
            return "html"
        # images are intentionally not supported without OCR
    if content_type:
        ct = content_type.lower()
        if "pdf" in ct:
            return "pdf"
        if "word" in ct or "officedocument" in ct:
            return "docx"
        if "html" in ct:
            return "html"
        if "text" in ct or "rtf" in ct:
            return "text"
    if filename and "." in filename:
        ext = filename.rsplit(".", 1)[1].lower()
        if ext in ("pdf",):
            return "pdf"
        if ext in ("docx", "doc"):
            return "docx"
        if ext in ("html", "htm"):
            return "html"
        if ext in ("txt", "rtf", "md"):
            return "text"
    return "text"


def extract_text_any(file_bytes: bytes, filename: Optional[str], content_type: Optional[str]) -> str:
    """
    Unified extraction without OCR (PDF/DOCX/HTML/TXT/RTF/MD).
    """
    kind = _guess_kind(file_bytes, filename, content_type)

    raw = ""
    try:
        if kind == "pdf":
            raw = extract_pdf_text_any(file_bytes)
        elif kind == "docx":
            raw = extract_docx_text(file_bytes)
        elif kind == "html":
            raw = extract_html_text(file_bytes)
        else:
            raw = extract_textlike_text(file_bytes)
    except Exception:
        raw = ""

    # Page-aware cleanup, then deep normalization
    pages = raw.split("\f") if "\f" in raw else raw.split("\n\n\n")
    pages = drop_repeated_header_footer(pages)
    return clean_text("\n\n".join(pages))
